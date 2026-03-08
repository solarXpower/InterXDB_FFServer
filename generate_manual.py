"""
Generate InterXDB_FFServer User Manual as a Word (.docx) document.
Run: py -3 generate_manual.py
Output: C:\AI Software\anviz\InterXDB_FFServer_User_Manual.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = r"C:\AI Software\anviz\InterXDB_FFServer_User_Manual.docx"

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x00, 0x33, 0x66)
MED_BLUE   = RGBColor(0x00, 0x56, 0x9B)
LIGHT_BLUE = RGBColor(0xD6, 0xE4, 0xF0)
ORANGE     = RGBColor(0xFF, 0x8C, 0x00)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY_BG    = RGBColor(0xF2, 0xF2, 0xF2)
BLACK      = RGBColor(0x00, 0x00, 0x00)

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helper: set paragraph shading ────────────────────────────────────────────
def shade_paragraph(para, hex_fill):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_fill)
    pPr.append(shd)

def shade_cell(cell, hex_fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_fill)
    tcPr.append(shd)

# ── Helper: add heading ───────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(16)
        run.bold = True
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = MED_BLUE
        run.font.size = Pt(13)
        run.bold = True
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = MED_BLUE
        run.font.size = Pt(11)
        run.bold = True
    return p

def body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.color.rgb = BLACK
    run.bold   = bold
    run.italic = italic
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def note_box(text):
    """Light-blue shaded note paragraph."""
    p = doc.add_paragraph()
    run = p.add_run("ℹ  " + text)
    run.font.size  = Pt(10.5)
    run.font.color.rgb = MED_BLUE
    run.italic = True
    shade_paragraph(p, "D6E4F0")
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

def warning_box(text):
    p = doc.add_paragraph()
    run = p.add_run("⚠  " + text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x99, 0x33, 0x00)
    run.bold = True
    shade_paragraph(p, "FFF3CC")
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

def two_col_table(rows, header=None):
    """Simple 2-column table with optional header row."""
    cols = 2
    tbl = doc.add_table(rows=len(rows) + (1 if header else 0), cols=cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Inches(2.2)
    tbl.columns[1].width = Inches(4.0)
    row_offset = 0
    if header:
        hrow = tbl.rows[0]
        for i, h in enumerate(header):
            cell = hrow.cells[i]
            shade_cell(cell, "005699")
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(11)
        row_offset = 1
    for r_idx, (col1, col2) in enumerate(rows):
        row = tbl.rows[r_idx + row_offset]
        bg = "F2F2F2" if r_idx % 2 == 0 else "FFFFFF"
        shade_cell(row.cells[0], bg)
        shade_cell(row.cells[1], bg)
        p0 = row.cells[0].paragraphs[0]
        run0 = p0.add_run(col1)
        run0.bold = True
        run0.font.size = Pt(10.5)
        p1 = row.cells[1].paragraphs[0]
        run1 = p1.add_run(col2)
        run1.font.size = Pt(10.5)
    return tbl

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
shade_paragraph(p, "003366")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\nInterXDB FFServer\n")
run.font.size  = Pt(30)
run.font.color.rgb = WHITE
run.bold = True
run2 = p.add_run("User Manual  •  Version 1.0\n\n")
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xFF)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run("Anviz VF30PRO / V300 Fingerprint Reader Management System\n"
               "OnSolar / The Next Software  •  March 2026")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
r.italic = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════
h1("Table of Contents")
toc = [
    ("1.", "Introduction"),
    ("2.", "System Requirements"),
    ("3.", "Installation"),
    ("4.", "First Launch & Connection"),
    ("5.", "Dashboard"),
    ("6.", "Punch Log"),
    ("7.", "Users & Enrollment"),
    ("8.", "Fingerprint Templates"),
    ("9.", "Network"),
    ("10.", "Settings"),
    ("11.", "Log File Format"),
    ("12.", "Troubleshooting"),
    ("13.", "Glossary"),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r1 = p.add_run(f"{num}  ")
    r1.bold = True
    r1.font.color.rgb = MED_BLUE
    r1.font.size = Pt(11)
    r2 = p.add_run(title)
    r2.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Introduction")
body("InterXDB FFServer is a Windows desktop application for managing Anviz VF30PRO and V300 "
     "series fingerprint access-control readers. It connects to one or more devices over a TCP/IP "
     "network, captures real-time attendance punches, manages enrolled users and fingerprint "
     "templates, and logs all activity to both a local SQLite database and a CrossChex-compatible "
     "CSV log file.")

doc.add_paragraph()
h2("Key Features")
features = [
    "Real-time punch capture with punch-type display (Check-In, Check-Out, Break-Out, Break-In, Overtime-Out, Overtime-In)",
    "Duplicate punch filtering — configurable time window prevents double entries",
    "CrossChex-compatible CSV log file (compatible with Anviz CrossChex software)",
    "SQLite database for persistent punch record storage",
    "User list viewer — shows all enrolled users with fingerprint count",
    "Delete users directly from the device",
    "Transfer users between multiple connected devices",
    "Download and save fingerprint templates to local files",
    "Upload fingerprint templates from local files back to a device",
    "Read and write device network configuration (IP, gateway, subnet, port)",
    "UDP network broadcast scan to discover Anviz devices on the LAN",
    "Supports both Client Mode (SDK connects to device) and Server Mode (device connects to SDK)",
    "Dark-theme professional GUI built with PyQt6",
]
for f in features:
    bullet(f)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════
h1("2. System Requirements")
two_col_table([
    ("Operating System",   "Windows 10 / 11  (64-bit only)"),
    ("Processor",          "Any 64-bit Intel or AMD processor"),
    ("RAM",                "Minimum 2 GB  (4 GB recommended)"),
    ("Disk Space",         "~150 MB for installation"),
    ("Network",            "Ethernet or Wi-Fi — same LAN as the Anviz device"),
    ("Device Compatibility", "Anviz VF30PRO, V300, and other TC_B_SDK compatible readers"),
    ("Python",             "Not required — the installer includes all dependencies"),
], header=["Component", "Requirement"])

doc.add_paragraph()
warning_box("InterXDB FFServer requires a 64-bit version of Windows. It will not run on 32-bit systems.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. INSTALLATION
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Installation")
h2("3.1  Running the Installer")
body("1.  Double-click  InterXDB_FFServer_Setup_v1.0.exe")
body("2.  If Windows SmartScreen appears, click  More info → Run anyway")
body("3.  Accept the User Account Control (UAC) prompt — administrator rights are required")
body("4.  Follow the on-screen wizard steps:")
bullet("Choose installation folder  (default: C:\\Program Files\\InterXDB FFServer\\)")
bullet("Choose whether to create a Desktop shortcut")
bullet("Click Install")
body("5.  Optionally tick  Launch InterXDB FFServer  and click  Finish")

doc.add_paragraph()
h2("3.2  Uninstalling")
body("Open  Settings → Apps  (or  Control Panel → Programs and Features), "
     "find  InterXDB FFServer 1.0, and click  Uninstall.  "
     "The uninstaller removes all program files and the Start Menu entry.  "
     "Your punch log (interxdb.db) and config file are also removed automatically.")

note_box("The application creates interxdb.db and interxdb_config.json in its installation folder "
         "on first run. Back up interxdb.db regularly to preserve punch records.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. FIRST LAUNCH & CONNECTION
# ══════════════════════════════════════════════════════════════════════════════
h1("4. First Launch & Connection")
h2("4.1  Launching the Application")
body("Launch InterXDB FFServer from the Start Menu or Desktop shortcut. "
     "The application opens in its dark-theme GUI and immediately attempts to start the SDK "
     "and connect to the device using the saved configuration.")

doc.add_paragraph()
h2("4.2  Connection Status Indicator")
body("The sidebar displays the current connection status:")
two_col_table([
    ("● Connected  (green)",    "SDK is running and at least one device is online"),
    ("● Disconnected  (red)",   "No device is connected"),
], header=["Indicator", "Meaning"])

doc.add_paragraph()
note_box("The status dot and device info (model, IP, firmware version, Machine ID) are shown "
         "at the bottom-left of the sidebar once a device connects.")

doc.add_paragraph()
h2("4.3  Configuring the Connection (Settings Panel)")
body("Before the first connection, go to Settings and enter your device details:")
two_col_table([
    ("Connection Mode",  "Client Mode — SDK connects TO the device (recommended for V300/VF30PRO in Server Mode)\n"
                         "Server Mode — Device connects TO the SDK (device must be in Client Mode)"),
    ("Device IP",        "IP address of the Anviz reader (e.g. 192.168.1.218)"),
    ("Device Port",      "TCP port of the device (default: 5010)"),
    ("Server Port",      "Local port for SDK Server Mode (default: 5010)"),
    ("Log Directory",    "Folder where the CrossChex CSV log file is saved"),
    ("Dedup Window (s)", "Seconds within which duplicate punches from the same user + punch type are ignored (default: 60)"),
], header=["Setting", "Description"])

doc.add_paragraph()
body("Click  Save Settings.  The application must be restarted for connection changes to take effect.")
warning_box("In Client Mode the device must be configured with a static IP address on the same "
            "LAN subnet as the PC running InterXDB FFServer.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Dashboard")
body("The Dashboard panel provides a live overview of connected devices and punch activity.")

h2("5.1  Stat Cards")
two_col_table([
    ("Devices",       "Number of Anviz readers currently connected to the SDK"),
    ("Today's Punches","Total punch events captured today (resets at midnight)"),
    ("Last Punch",    "Timestamp and User ID of the most recent punch event"),
    ("Log File",      "Status of the CrossChex CSV log file (path shown)"),
], header=["Card", "Description"])

doc.add_paragraph()
h2("5.2  Connected Devices Table")
body("The table lists each connected device with its index, Machine ID, model name, "
     "and network address (IP:Port).  Devices appear automatically when they connect and "
     "are removed when they disconnect.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. PUNCH LOG
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Punch Log")
body("The Punch Log panel shows all punch events captured since the application started, "
     "in real time.  Each row represents one punch event.")

h2("6.1  Punch Log Columns")
two_col_table([
    ("User ID",    "Numeric employee ID enrolled on the device"),
    ("Code",       "Short punch-type code: I, O, BO, BI, LO, LI"),
    ("Description","Full punch-type name (see table below)"),
    ("Verify",     "Verification method used: FP (fingerprint), Card, PIN, etc."),
    ("Timestamp",  "Date and time of the punch event (YYYY-MM-DD HH:MM:SS)"),
    ("Device",     "Device index that recorded the punch"),
    ("Source",     "RECORD (stored log) or LIVE (real-time attendance event)"),
], header=["Column", "Description"])

doc.add_paragraph()
h2("6.2  Punch Type Codes")
two_col_table([
    ("I",  "Check-In  — standard daily clock-in"),
    ("O",  "Check-Out — standard daily clock-out"),
    ("BO", "Break-Out — leaving for a break"),
    ("BI", "Break-In  — returning from a break"),
    ("LO", "Overtime-Out — leaving after overtime"),
    ("LI", "Overtime-In  — starting overtime"),
], header=["Code", "Description"])

doc.add_paragraph()
h2("6.3  Export to CSV")
body("Click  Export CSV  to save the current punch log table to a comma-separated file. "
     "A file-save dialog will prompt for the destination folder and filename.")

doc.add_paragraph()
h2("6.4  Duplicate Punch Filtering")
body("InterXDB FFServer automatically ignores duplicate punches. A punch is considered a "
     "duplicate if the same User ID submits the same punch type within the configured "
     "Dedup Window (default 60 seconds). The duplicate is silently discarded — it does not "
     "appear in the log or database.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. USERS & ENROLLMENT
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Users & Enrollment")
body("The Users panel lists all users enrolled on a connected device and provides tools to "
     "delete or transfer users.")

h2("7.1  Loading the User List")
body("1.  Select the target device from the  Device  drop-down.")
body("2.  Click  Load Users.")
body("3.  The table populates with all enrolled users. A progress bar is shown during loading.")

doc.add_paragraph()
h2("7.2  User Table Columns")
two_col_table([
    ("User ID",       "Numeric employee ID"),
    ("Name",          "Employee name stored on the device (if set)"),
    ("Card ID",       "Card number enrolled for this user (0xFFFFFFFF = no card)"),
    ("FP Templates",  "Number of fingerprint templates enrolled (bits 0–9 of Fp_Status bitmask)"),
    ("Dept",          "Department ID assigned to the user"),
], header=["Column", "Description"])

doc.add_paragraph()
h2("7.3  Deleting Users")
body("1.  Select one or more rows in the user table (hold Ctrl or Shift to multi-select).")
body("2.  Click  Delete Selected.")
body("3.  Confirm the deletion in the dialog box.")
warning_box("Deleting a user permanently removes them from the device, including all their "
            "fingerprint templates and card assignments. This action cannot be undone.")

doc.add_paragraph()
h2("7.4  Transferring Users Between Devices")
body("1.  Load users from the source device.")
body("2.  Select the users to transfer.")
body("3.  Click  Transfer Selected →.")
body("4.  Choose the target device from the drop-down list.")
body("5.  Click OK. The selected user records are copied to the target device.")
note_box("Transfer copies user information (ID, name, card) but does NOT copy fingerprint templates. "
         "Use the Fingerprints panel to download and re-upload fingerprint data separately.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. FINGERPRINT TEMPLATES
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Fingerprint Templates")
body("The Fingerprints panel allows you to download fingerprint templates from a device to "
     "local files, and to upload previously saved templates back to a device. "
     "This is useful for backup, migration, and cloning devices.")

h2("8.1  Download All Fingerprints")
body("1.  Go to the  Users panel  and click  Load Users  first (required to enumerate user IDs).")
body("2.  Switch to the  Fingerprints panel.")
body("3.  Click  Download All Fingerprints.")
body("4.  The application iterates over every enrolled user and every finger index (0–9), "
     "downloading any template that exists on the device.")
body("5.  Templates are saved as binary files in the configured Log Directory:")
bullet("Filename format:  <UserID>_finger<Index>.fp")
bullet("Example:  900_finger0.fp,  900_finger1.fp")
body("6.  The progress log on-screen reports each download result.")

doc.add_paragraph()
h2("8.2  Upload All Fingerprints from Folder")
body("1.  Ensure the target users already exist on the device (create them first if needed).")
body("2.  Click  Upload All Fingerprints from Folder.")
body("3.  Select the folder containing the .fp files.")
body("4.  The application uploads every .fp file found, matching filename patterns "
     "<UserID>_finger<Index>.fp  to the correct user and finger slot on the device.")

doc.add_paragraph()
note_box("Fingerprint templates are device-specific binary data. Templates downloaded from a "
         "VF30PRO can generally be uploaded to another VF30PRO unit of the same firmware generation.")

warning_box("Do not edit .fp files manually — they contain raw biometric binary data that will be "
            "corrupted by any text modification.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. NETWORK
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Network")
body("The Network panel lets you discover Anviz devices on the LAN via UDP broadcast, "
     "and read or write the network configuration of a connected device.")

h2("9.1  UDP Device Scan")
body("Click  Scan Network  to broadcast a UDP discovery packet. All Anviz devices on the "
     "local network that respond are listed in the Discovered Devices table, showing their "
     "IP address, MAC address, device model, and current port.")
note_box("The UDP scan only finds devices on the same broadcast domain (same LAN subnet). "
         "Devices behind a router on a different subnet will not appear.")

doc.add_paragraph()
h2("9.2  Reading Network Configuration")
body("1.  Select a connected device from the  Device  drop-down.")
body("2.  Click  Read Config.")
body("3.  The current device settings populate the form fields.")

doc.add_paragraph()
h2("9.3  Writing Network Configuration")
two_col_table([
    ("IP Address",   "New static IP for the device"),
    ("Subnet Mask",  "Network subnet mask (e.g. 255.255.255.0)"),
    ("Gateway",      "Default gateway IP"),
    ("DNS",          "DNS server IP"),
    ("Port",         "TCP port the device listens on (default: 5010)"),
    ("Net Mode",     "0 = Server Mode (device acts as TCP server)\n1 = Client Mode (device connects to a host)"),
], header=["Field", "Description"])

doc.add_paragraph()
body("After entering new values, click  Write Config  to apply them to the device.")
warning_box("Changing the device IP address will disconnect it immediately. Update the IP in "
            "Settings and restart the application to reconnect.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. SETTINGS
# ══════════════════════════════════════════════════════════════════════════════
h1("10. Settings")
body("The Settings panel stores all application configuration. Changes are saved to "
     "interxdb_config.json in the installation folder.")

two_col_table([
    ("Connection Mode",  "client — SDK connects to the device IP:Port\n"
                         "server — SDK listens; device connects to the PC"),
    ("Device IP",        "IP address of the Anviz device"),
    ("Device Port",      "TCP port of the Anviz device (default 5010)"),
    ("Server Port",      "Local TCP port for SDK Server Mode (default 5010)"),
    ("Log Directory",    "Folder path for the CrossChex CSV log file\n"
                         "(can be a local folder or a mapped network drive)"),
    ("Dedup Window (s)", "Duplicate punch suppression window in seconds (default 60)\n"
                         "Set to 0 to disable duplicate filtering"),
], header=["Setting", "Description"])

doc.add_paragraph()
body("Click  Save Settings  to write changes to disk. "
     "Restart the application after changing any connection setting.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. LOG FILE FORMAT
# ══════════════════════════════════════════════════════════════════════════════
h1("11. Log File Format")
body("InterXDB FFServer writes punch records to a CrossChex-compatible CSV log file. "
     "The file is named:  Punchtypes format <Device IP>.log")
body("Example filename:  Punchtypes format 192.168.1.218.log", bold=True)

doc.add_paragraph()
h2("11.1  File Structure")
body("The first line is a fixed header:")
p = doc.add_paragraph()
shade_paragraph(p, "F2F2F2")
p.add_run("C,GetAllLogs,SUCCESSFUL").font.name = "Courier New"

body("Each subsequent line is one punch record:")
p2 = doc.add_paragraph()
shade_paragraph(p2, "F2F2F2")
p2.add_run('UserID,PunchCode,"YYYY-MM-DD HH:MM:SS",0,0').font.name = "Courier New"

doc.add_paragraph()
h2("11.2  Example")
examples = [
    "C,GetAllLogs,SUCCESSFUL",
    '900,I,"2026-03-07 20:22:56",0,0',
    '900,O,"2026-03-07 20:23:10",0,0',
    '255,I,"2026-03-07 20:43:42",0,0',
]
for ex in examples:
    p = doc.add_paragraph()
    shade_paragraph(p, "F2F2F2")
    run = p.add_run(ex)
    run.font.name = "Courier New"
    run.font.size = Pt(10)

doc.add_paragraph()
note_box("This format is directly compatible with Anviz CrossChex Software for import and reporting.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12. TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════════════════
h1("12. Troubleshooting")

issues = [
    (
        "Failed to start SDK — DLL not found",
        "The tc-b_new_sdk.dll file is missing from the _internal folder. "
        "Re-install the application using the setup installer."
    ),
    (
        "Status shows Disconnected after launch",
        "Check that:\n"
        "• The device IP and port in Settings are correct\n"
        "• The device is powered on and connected to the network\n"
        "• Windows Firewall is not blocking the connection (allow port 5010)\n"
        "• The PC and device are on the same LAN subnet"
    ),
    (
        "Punch events not appearing in real time",
        "Ensure the device is in Server Mode (Net Mode = 0) when using Client Mode in InterXDB FFServer. "
        "After connecting, place your finger on the reader — punch events appear within 1–2 seconds."
    ),
    (
        "FP Templates shows 0 for all users",
        "The device may not populate the Fp_Status field in the list response. "
        "Use the Fingerprints panel to perform an actual download to verify enrolled templates."
    ),
    (
        "Log file not visible in the folder",
        "Ensure Windows is set to show file extensions (View → Show → File name extensions). "
        "The log file has the .log extension and is created in the configured Log Directory."
    ),
    (
        "Application icon shows old image in Explorer",
        "Right-click the EXE → Properties → OK to refresh the Windows icon cache."
    ),
    (
        "Failed to load Python DLL error",
        "The EXE was separated from the _internal folder. "
        "Always keep InterXDB_FFServer.exe and the _internal folder together in the same directory."
    ),
    (
        "Settings reset after update",
        "interxdb_config.json is stored in the installation folder. "
        "If it is deleted or the folder permissions change, settings revert to defaults."
    ),
]

for problem, solution in issues:
    h3("Problem: " + problem)
    body("Solution: " + solution)
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 13. GLOSSARY
# ══════════════════════════════════════════════════════════════════════════════
h1("13. Glossary")
two_col_table([
    ("SDK",           "Software Development Kit — the Anviz TC_B_SDK_V64 library (tc-b_new_sdk.dll) "
                      "used to communicate with the device"),
    ("Client Mode",   "Connection mode where the InterXDB FFServer SDK actively connects to the device "
                      "IP address and port"),
    ("Server Mode",   "Connection mode where the SDK listens on a local port and the device initiates "
                      "the connection to the PC"),
    ("VF30PRO",       "Anviz fingerprint reader model — optical sensor, supports up to 10 fingerprints "
                      "per user, TypeFlag 0x02101204"),
    ("V300",          "Anviz access-control reader compatible with the TC_B_SDK"),
    ("MachineId",     "Unique numeric identifier assigned to each Anviz device"),
    ("Fp_Status",     "32-bit bitmask where bits 0–9 indicate whether finger indices 0–9 are enrolled"),
    ("DevIdx",        "Device index (1-based) assigned by the SDK to each connected device"),
    ("CrossChex",     "Anviz attendance management software — InterXDB FFServer uses the same log "
                      "format for compatibility"),
    ("UDP Scan",      "Network broadcast discovery using CCHex_Udp_Search_Dev — finds Anviz devices "
                      "on the local LAN"),
    ("Dedup Window",  "Time window (in seconds) during which repeated punches of the same type by the "
                      "same user are suppressed"),
    (".fp File",      "Binary file containing a raw fingerprint template downloaded from the device"),
], header=["Term", "Definition"])

doc.add_paragraph()

# ── Footer note ───────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
shade_paragraph(p, "003366")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("InterXDB FFServer v1.0  •  User Manual  •  © 2026 OnSolar / The Next Software  •  All rights reserved")
run.font.size = Pt(9)
run.font.color.rgb = WHITE

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Manual saved to: {OUTPUT}")
